import pandas as pd
from docx import Document
from docx.shared import RGBColor
address = input('输入文件地址')
Num = input('输入需要判读的行数')
finalDoc = input('输入需要生成的文档名称')
excel = pd.read_excel(address + '.xlsx')
Samples = excel.loc[:int(Num)-1, 'Sample']
aPositive = excel.loc[excel['95d'] > 9.47, ['Sample', '95d']]
bPositive = excel.loc[excel['94E'] > 19.5, ['Sample', '94E']]
cPositive = excel.loc[excel['132A'] > 12, ['Sample', '132A']]
dPositive = excel.loc[excel['90F'] > 18, ['Sample', '90F']]
ePositive = excel.loc[excel[8601] > 6.28, ['Sample', 8601]]
fPositive = excel.loc[excel[2801] > 9.93, ['Sample', 2801]]
gPositive = excel.loc[excel[4203] > 6.97, ['Sample', 4203]]

title = ['SampleID:', '95d:', '94E:', '132A:', '90F:', '8601:', '2801:', '4203:', '判断:']
doc = Document()
table = doc.add_table(int(Num) + 1, 9, 'Table Grid')

A = 0
while A < 9:
    table.rows[0].cells[A].add_paragraph(title[A])
    A += 1
for column in table.columns:
  for cell in column.cells:
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(250,0,0)
            run.font.bold = True

for index in Samples:
    IDs = Samples.values
B = 1
for ID in IDs:
    a = '阴性'
    b = '阴性'
    c = '阴性'
    d = '阴性'
    e = '阴性'
    f = '阴性'
    g = '阴性'
    for i in aPositive['Sample']:
        if ID == i:
            a = '阳性'
    for i in bPositive['Sample']:
        if ID == i:
            b = '阳性'
    for i in cPositive['Sample']:
        if ID == i:
            c = '阳性'
    for i in dPositive['Sample']:
        if ID == i:
            d = '阳性'
    for i in ePositive['Sample']:
        if ID == i:
            e = '阳性'
    for i in fPositive['Sample']:
        if ID == i:
            f = '阳性'
    for i in gPositive['Sample']:
        if ID == i:
            g = '阳性'
    results = [a, b, c, d, e, f, g]
    y = 0
    for s in results:
        if s == '阳性':
            y += 1
    x = 0
    if y < 1:
        x = '阴性'
    elif y == 1:
        x = '低风险'
    elif y == 2:
        x = '中风险'
    else:
        x = '高风险'
    patient = {'SampleID:': ID, '95d:': a, '94E:': b, '132A:': c,
               '90F:': d, '8601:': e, '2801:': f, '4203:': g, '判断:': x}
    TEXT = [ID, a, b, c, d, e, f, g, x]
    A = 0
    while A < 9:
        table.rows[B].cells[A].add_paragraph(TEXT[A])
        A += 1
    B += 1

doc.save(str(finalDoc) + '.docx')
