# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import  Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import wx

doc = Document()

tab =doc.add_table(rows=1,cols=1)   
cell=tab.cell(0,0)  
c_p1 =cell.paragraphs[0]
c_p1.paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
c_run1=c_p1.add_run()
c_run1.add_picture('xiansai.png',width=Cm(2))

c_p2 =cell.paragraphs[0]
c_p2.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
c_run2=c_p2.add_run('AD综合会诊专用申请单')
c_run2.font.size = Pt(20)
c_run2.font.bold = True
c_run2.font.color.rgb = RGBColor(0,134,139)

c_p3 =cell.paragraphs[0]
c_p3.paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
c_run3=c_p3.add_run()
c_run3.add_picture('tiaoma.png', width=Cm(5))


b = doc.add_paragraph()
run = b.add_run('患者基本就诊信息：')
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = RGBColor(75,0,130)

table = doc.add_table(rows=3, cols=4, style ='Table Grid')
table.cell(0,0).text = '送检医院：'
table.cell(1,0).text = '患者姓名:'
table.cell(2,0).text = '标本类型：'
for column in table.columns:
  for cell in column.cells:
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(250,0,0)
            run.font.bold = True
table.cell(0,1).text = '门诊/住院号：'
table.cell(0,2).text = '送诊医生：'
table.cell(0,3).text = '医生电话：'
table.cell(1,1).text = '性别:  男      女'

table.cell(1,2).text = '年龄：'
table.cell(1,3).text = '科室：'
table.cell(2,1).text = '血清         全血'
table.cell(2,2).text = '取材时间：   月    日'
table.cell(2,3).text = '送检时间：   月    日'

table1 = doc.add_table(rows=4, cols=3, style ='Table Grid')
table1.cell(0,0).text = '临床诊断：'
d = table1.cell(0,0).text
table1.cell(0,0).merge(table.cell(2,0))
table1.cell(0,1).text = '是否接受过免疫治疗'
table1.cell(1,1).text = '是否最近接受免疫调节剂的治疗'
table1.cell(2,1).text = '是否有风湿免疫性疾病'
table1.cell(3,1).text = '是否存在肿瘤史'
for column in table1.columns:
  for cell in column.cells:
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(250,0,0)
            run.font.bold = True
table1.cell(0,2).text = '是     否'
table1.cell(1,2).text = '是     否'
table1.cell(2,2).text = '是     否'
table1.cell(3,2).text = '是     否'

table2 = doc.add_table(1,2,'Table Grid')
table2.cell(0,0).text = '主要临床表现：'
table2.cell(0,0).width = Cm(4)
table2.cell(0,1).text = '起病形式： 急性      亚急性       慢性           病程：      日/月/年'

table3 = doc.add_table(1,1,'Table Grid')
table3.cell(0,0).text = '短期记忆力下降      睡眠障碍      认知功能异常     味觉下降      嗅觉下降      言语障碍\n' \
                        '意识模糊   食欲减退    自主神经功能   共济失调    其他'

table4 = doc.add_table(3,2,'Table Grid')

table4.cell(0,0).text = '历史摘要:'
table4.cell(1,1).text = '高血压   高血脂   '
table4.cell(2,0).text = 'IDDD:'
table4.cell(2,1).text = '正常   其他描述'


table5 = doc.add_table(1,1,'Table Grid')
table5.cell(0,0).text = '\n\n'

b = doc.add_paragraph()
run = b.add_run('标本送检要求：')
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = RGBColor(75,0,130)

table6 = doc.add_table(2,3,'Table Grid')
table6.cell(0,0).text = '项目明细'
table6.cell(0,1).text = '标本类型'
table6.cell(0,2).text = '收费'
for row in table6.rows:
  for cell in row.cells:
      paragraphs = cell.paragraphs
      for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(250,0,0)
            run.font.bold = True
table6.cell(1,0).text = '阿尔茨海默-七项自身抗体检测（Elisa）'
table6.cell(1,1).text = '血清（不少于1ml）'
table6.cell(1,2).text = '4000元/人（次）'


table7 = doc.add_table(1,1,'Table Grid')
table7.cell(0,0).text = '\n样本要求：\n'+'1、样本为外周静脉血血清，每份血清体积不小于1ml；未大量饮水；\n' \
                                    '2、血样采集后，应不加抗凝剂，并放置在干净、干燥的试管内2～8℃静置1小时，离心（3000rpm×8min），吸取血清；\n' \
                                    '3、血清样本室温下存放不要超过8小时；血清样本如在8小时内不检测，应在2～8℃下保存，但时间不超过48小时；超过48小时或需要运输，应在-20℃以下保存，避免反复冻融（冻融2次以上不合格）。冷冻的样本使用前要融化后混匀；\n' \
                                    '4、收到样本后3个工作日出报告；\n' \
                                    '5、本项目来自全球最新的前沿科技，拥有全球专利，处于注册阶段；现主要用于科研和临床研究；\n' \
                                    '6、收费4000元/人（次）包括阿尔兹海默七个抗体的检测。'

doc.save('new.docx')
