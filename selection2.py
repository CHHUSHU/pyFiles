import pandas as pd
from docx import Document

excel = pd.read_excel('C:\\Users\\xiaoy\\Documents\\pyFiles\\test1.xlsx')
aPositive = excel.loc[excel['95d']>9.47,['Sample','95d']]
bPositive = excel.loc[excel['94E']>19.5,['Sample','94E']]
cPositive = excel.loc[excel['132A']>12, ['Sample','132A']]
dPositive = excel.loc[excel['90F']>18, ['Sample','90F']]
ePositive = excel.loc[excel[8601]>6.28, ['Sample',8601]]
fPositive = excel.loc[excel[2801]>9.93, ['Sample',2801]]
gPositive = excel.loc[excel[4203]>6.97, ['Sample',4203]]

name = input('Enter a patient name\n')
ID = input('Enter a sample ID\n')
a = 'negative'
b = 'negative'
c = 'negative'
d = 'negative'
e = 'negative'
f = 'negative'
g = 'negative'
for i in aPositive['Sample']:
    if ID == i:
        a = 'positive'    
for i in bPositive['Sample']:
    if ID == i:
        b = 'positive'
for i in cPositive['Sample']:
    if ID == i:
        c = 'positive'
for i in dPositive['Sample']:
    if ID == i:
        d = 'positive'
for i in ePositive['Sample']:
    if ID == i:
        e = 'positive'
for i in fPositive['Sample']:
    if ID == i:
        f = 'positive'
for i in gPositive['Sample']:
    if ID == i:
        g= 'positive'

patient = {'Name:': name, 'SampleID:': ID, '95d:':a, '94E:':b, '132A:':c,
           '90F:': d, '8601:': e, '2801:': f, '4203:': g}
print(patient)
results = [a,b,c,d,e,f,g]
y = 0
for s in results:
    if s == 'positive':
        y += 1
print(y)

x = 0
if y < 1:
    print('阴性')
    x = '阴性'
elif y == 1:
    print('低风险')
    x = '低风险'
elif y == 2:
    print('中风险')
    x = '中风险'
else:
    print('高风险')
    x = '高风险'



doc = Document()
table = doc.add_table(2,9)
table.cell(0,0).text = '患者姓名:'
table.cell(1,0).text = name
table.cell(0,1).text = '血清样本：'
table.cell(1,1).text = ID
table.cell(0,2).text = '95d:'
table.cell(1,2).text = a
table.cell(0,3).text = '94E:'
table.cell(1,3).text = b
table.cell(0,4).text = '132A:'
table.cell(1,4).text = c
table.cell(0,5).text = '90F:'
table.cell(1,5).text = d
table.cell(0,6).text = '8601:'
table.cell(1,6).text = e
table.cell(0,7).text = '2801:'
table.cell(1,7).text = f
table.cell(0,8).text = '4203:'
table.cell(1,8).text = g

table1 = doc.add_table(1,1)
table1.cell(0,0).text = x
doc.save('result.docx')