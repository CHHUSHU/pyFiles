import pandas as pd
from docx import Document

excel = pd.read_excel('C:\\Users\\xiaoy\\Documents\\pyFiles\\test1.xlsx')
Samples = excel.loc[:40, 'Sample']
aPositive = excel.loc[excel['95d']>9.47,['Sample','95d']]
bPositive = excel.loc[excel['94E']>19.5,['Sample','94E']]
cPositive = excel.loc[excel['132A']>12, ['Sample','132A']]
dPositive = excel.loc[excel['90F']>18, ['Sample','90F']]
ePositive = excel.loc[excel[8601]>6.28, ['Sample',8601]]
fPositive = excel.loc[excel[2801]>9.93, ['Sample',2801]]
gPositive = excel.loc[excel[4203]>6.97, ['Sample',4203]]


title = ['SampleID:','95d:','94E:','132A:','90F:','8601:','2801:','4203:','判断:']
doc = Document()
table = doc.add_table(42,9,'Table Grid')

A = 0
while A<9: 
    table.rows[0].cells[A].add_paragraph(title[A])
    A+=1

for index in Samples:
    IDs = Samples.values
B = 1
for ID in IDs:
    a = 'negative'
    b = 'negative'
    c = 'negative'
    d = 'negative'
    e = 'negative'
    f = 'negative'
    g = 'negative'
    for i in aPositive['Sample']:
        if ID == i:
            a = 'positive'
    for i in bPositive['Sample']:
        if ID == i:
            b = 'positive'
    for i in cPositive['Sample']:
        if ID == i:
            c = 'positive'
    for i in dPositive['Sample']:
        if ID == i:
            d = 'positive'
    for i in ePositive['Sample']:
        if ID == i:
            e = 'positive'
    for i in fPositive['Sample']:
        if ID == i:
            f = 'positive'
    for i in gPositive['Sample']:
        if ID == i:
            g= 'positive'
    results = [a,b,c,d,e,f,g]
    y = 0
    for s in results:
        if s == 'positive':
            y += 1
    print(y)
    results = [a,b,c,d,e,f,g]
    y = 0
    for s in results:
        if s == 'positive':
            y += 1
    print(y)
    x = 0
    if y < 1:
        print('阴性')
        x = '阴性'
    elif y == 1:
        print('低风险')
        x = '低风险'
    elif y == 2:
        print('中风险')
        x = '中风险'
    else:
        print('高风险')
        x = '高风险'
    patient = {'SampleID:': ID, '95d:': a, '94E:': b, '132A:': c,
                   '90F:': d, '8601:': e, '2801:': f, '4203:': g, '判断:':x}
    print(patient)

    TEXT = [ID,a,b,c,d,e,f,g,x]
    A = 0
    while A<9: 
        table.rows[B].cells[A].add_paragraph(TEXT[A])
        A+=1
    B+=1


doc.save('Test.docx')
